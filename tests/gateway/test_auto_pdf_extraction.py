"""Gateway-side automatic PDF extraction (GatewayRunner._auto_extract_pdf).

Reliable, model-independent PDF handling wired into the inbound pipeline:
  - text-layer PDFs are read with pymupdf and inlined into the message,
  - scanned PDFs (no text layer) fall back to render + the vision auxiliary,
  - non-PDF files are ignored (returns None),
  - pymupdf being unavailable degrades gracefully (returns None).

pymupdf is not a test-env dependency, so we inject a fake module.
"""

from __future__ import annotations

import asyncio
import sys
import types

from unittest.mock import AsyncMock

import pytest

from gateway.run import GatewayRunner


def _fake_pymupdf(pages_text):
    m = types.ModuleType("pymupdf")

    class _Page:
        def __init__(self, t):
            self._t = t

        def get_text(self):
            return self._t

    class _Doc:
        def __init__(self, texts):
            self._pages = [_Page(t) for t in texts]
            self.page_count = len(texts)

        def __iter__(self):
            return iter(self._pages)

        def __getitem__(self, i):
            return self._pages[i]

        def close(self):
            pass

    m.open = lambda path: _Doc(pages_text)
    return m


class _Self:
    """Minimal stand-in for GatewayRunner (methods only touch self._vision_...)."""


def test_non_pdf_returns_none(tmp_path):
    p = tmp_path / "x.bin"
    p.write_bytes(b"NOT-A-PDF-HEADER........")
    r = asyncio.run(GatewayRunner._auto_extract_pdf(_Self(), str(p), "x.bin"))
    assert r is None


def test_missing_file_returns_none(tmp_path):
    r = asyncio.run(GatewayRunner._auto_extract_pdf(_Self(), str(tmp_path / "nope.pdf"), "nope.pdf"))
    assert r is None


def test_text_layer_pdf_is_inlined(tmp_path, monkeypatch):
    p = tmp_path / "r.pdf"
    p.write_bytes(b"%PDF-1.4\n" + b"0" * 32)  # magic bytes ok; content irrelevant (fake pymupdf)
    monkeypatch.setitem(
        sys.modules, "pymupdf", _fake_pymupdf(["Aryaduta Bali\nTotal USD 336.35\nHSIN-YI LIU"])
    )
    r = asyncio.run(GatewayRunner._auto_extract_pdf(_Self(), str(p), "receipt.pdf"))
    assert r is not None
    assert "Auto-extracted text" in r
    assert "Aryaduta Bali" in r and "336.35" in r  # inlined verbatim, no model/tool needed


def test_scanned_pdf_falls_back_to_vision(tmp_path, monkeypatch):
    p = tmp_path / "s.pdf"
    p.write_bytes(b"%PDF-1.4\n" + b"0" * 32)
    monkeypatch.setitem(sys.modules, "pymupdf", _fake_pymupdf([""]))  # empty text -> scanned
    s = _Self()
    s._vision_read_scanned_pdf = AsyncMock(return_value="[Auto-read scanned PDF] Hotel Foo, TWD 10,826")
    r = asyncio.run(GatewayRunner._auto_extract_pdf(s, str(p), "scan.pdf"))
    assert r is not None and "Hotel Foo" in r
    s._vision_read_scanned_pdf.assert_awaited_once()


def test_pymupdf_unavailable_degrades(tmp_path, monkeypatch):
    p = tmp_path / "r.pdf"
    p.write_bytes(b"%PDF-1.4\n" + b"0" * 32)
    # Force `import pymupdf` to fail.
    monkeypatch.setitem(sys.modules, "pymupdf", None)
    r = asyncio.run(GatewayRunner._auto_extract_pdf(_Self(), str(p), "r.pdf"))
    assert r is None


def test_vision_read_renders_and_transcribes(tmp_path, monkeypatch):
    # Fake a scanned doc + vision endpoint end-to-end for _vision_read_scanned_pdf.
    import gateway.run as run_mod

    class _Pix:
        def save(self, path):
            with open(path, "wb") as f:
                f.write(b"\x89PNG\r\n\x1a\n")

    class _Page:
        def get_pixmap(self, dpi=170):
            return _Pix()

    class _Doc:
        page_count = 1

        def __getitem__(self, i):
            return _Page()

    import tools.vision_tools as vt
    monkeypatch.setattr(
        vt, "vision_analyze_tool",
        AsyncMock(return_value='{"success": true, "analysis": "Aryaduta Bali  USD 336.35"}'),
    )
    r = asyncio.run(GatewayRunner._vision_read_scanned_pdf(_Self(), _Doc(), "scan.pdf", 8))
    assert r is not None
    assert "Aryaduta Bali" in r and "Page 1" in r


# ---------------------------------------------------------------------------
# _auto_extract_document — generic dispatcher (text / csv / docx / xlsx / pdf)
# ---------------------------------------------------------------------------

def test_doc_text_file_inlined(tmp_path):
    p = tmp_path / "notes.txt"
    p.write_text("Aryaduta Bali\nTotal USD 336.35\n", encoding="utf-8")
    r = asyncio.run(GatewayRunner._auto_extract_document(_Self(), str(p), "text/plain", "notes.txt"))
    assert r and "Auto-extracted text" in r and "Aryaduta Bali" in r


def test_doc_csv_inlined(tmp_path):
    p = tmp_path / "data.csv"
    p.write_text("name,amount\nAryaduta Bali,336.35\n", encoding="utf-8")
    r = asyncio.run(GatewayRunner._auto_extract_document(_Self(), str(p), "text/csv", "data.csv"))
    assert r and "336.35" in r


def test_doc_docx_inlined(tmp_path):
    import docx
    d = docx.Document()
    d.add_paragraph("Aryaduta Bali")
    d.add_paragraph("Total USD 336.35")
    p = tmp_path / "receipt.docx"
    d.save(str(p))
    r = asyncio.run(GatewayRunner._auto_extract_document(_Self(), str(p), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "receipt.docx"))
    assert r and "Word document text" in r and "Aryaduta Bali" in r and "336.35" in r


def test_doc_xlsx_inlined(tmp_path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["hotel", "amount"])
    ws.append(["Aryaduta Bali", 336.35])
    p = tmp_path / "receipt.xlsx"
    wb.save(str(p))
    r = asyncio.run(GatewayRunner._auto_extract_document(_Self(), str(p), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "receipt.xlsx"))
    assert r and "spreadsheet" in r and "Aryaduta Bali" in r and "336.35" in r


def test_doc_unsupported_returns_none(tmp_path):
    # pptx / legacy office / archive -> no bundled extractor -> None (context note fallback)
    p = tmp_path / "deck.pptx"
    p.write_bytes(b"PK\x03\x04" + b"0" * 64)  # zip-based, but no pptx extractor
    r = asyncio.run(GatewayRunner._auto_extract_document(_Self(), str(p), "application/vnd.openxmlformats-officedocument.presentationml.presentation", "deck.pptx"))
    assert r is None


def test_doc_pdf_delegates(tmp_path, monkeypatch):
    p = tmp_path / "r.pdf"
    p.write_bytes(b"%PDF-1.4\n" + b"0" * 32)
    monkeypatch.setitem(sys.modules, "pymupdf", _fake_pymupdf(["Aryaduta Bali USD 336.35"]))
    s = _Self()
    s._auto_extract_pdf = GatewayRunner._auto_extract_pdf.__get__(s, _Self)  # real PDF branch
    r = asyncio.run(GatewayRunner._auto_extract_document(s, str(p), "application/pdf", "r.pdf"))
    assert r and "Aryaduta Bali" in r
